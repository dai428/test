"""基于NSGA-II的LCVR Stokes偏振计多目标优化论文 V5.1 — 核心期刊格式"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def srf(run, cn='宋体', en='Times New Roman', sz=Pt(12), b=False):
    run.font.name = en; run.font.size = sz; run.bold = b
    run.element.rPr.rFonts.set(qn('w:eastAsia'), cn)

def h1(doc, txt):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_before = Pt(12); pf.space_after = Pt(6); pf.line_spacing = 1.5
    srf(p.add_run(txt), cn='黑体', en='黑体', sz=Pt(16), b=True)

def h2(doc, txt):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format; pf.space_before = Pt(8); pf.space_after = Pt(4); pf.line_spacing = 1.5
    srf(p.add_run(txt), cn='黑体', en='黑体', sz=Pt(14), b=True)

def h3(doc, txt):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format; pf.space_before = Pt(6); pf.space_after = Pt(3); pf.line_spacing = 1.5
    srf(p.add_run(txt), cn='黑体', en='黑体', sz=Pt(12), b=True)

def body(doc, txt, indent=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    if indent: pf.first_line_indent = Pt(24)
    pf.line_spacing = 1.5
    srf(p.add_run(txt))

def ctr(doc, txt, sz=Pt(12)):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.line_spacing = 1.5
    srf(p.add_run(txt), sz=sz)

def tbl(doc, hdrs, rows, cap=None):
    if cap:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format; pf.space_before = Pt(6); pf.space_after = Pt(3)
        srf(p.add_run(cap), cn='黑体', sz=Pt(10.5), b=True)
    t = doc.add_table(rows=1, cols=len(hdrs))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i, h in enumerate(hdrs):
        c = t.rows[0].cells[i]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        srf(p.add_run(h), cn='黑体', sz=Pt(10.5), b=True)
    for rd in rows:
        row_cells = t.add_row().cells
        for i, v in enumerate(rd):
            row_cells[i].text = ''
            p = row_cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            srf(p.add_run(v), sz=Pt(10.5))
    doc.add_paragraph()

def figcap(doc, txt):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_before = Pt(3); pf.space_after = Pt(6)
    srf(p.add_run(txt), sz=Pt(10.5))

def cover(doc, title, items):
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srf(p.add_run(title), cn='黑体', en='黑体', sz=Pt(22), b=True)
    doc.add_paragraph(); doc.add_paragraph()
    for l, v in items:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        srf(p.add_run(f'{l}：{v}'), sz=Pt(14))

def abst(doc, txt, keys):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_before = Pt(12); pf.space_after = Pt(6)
    srf(p.add_run('摘  要'), cn='黑体', en='黑体', sz=Pt(16), b=True)
    doc.add_paragraph(); body(doc, txt)
    p = doc.add_paragraph(); pf = p.paragraph_format; pf.first_line_indent = Pt(24); pf.line_spacing = 1.5
    srf(p.add_run('关键词：'), cn='黑体', sz=Pt(12), b=True)
    srf(p.add_run(keys))

def refs(doc, r):
    h1(doc, '参考文献')
    for s in r:
        p = doc.add_paragraph(); pf = p.paragraph_format; pf.line_spacing = 1.5
        srf(p.add_run(s), sz=Pt(10.5))

def gen(doc):
    # ═══ 封面 ═══
    cover(doc, '基于NSGA-II的LCVR Stokes偏振计\n多目标优化', [('作者', '（待填）'), ('单位', '（待填）')])
    doc.add_page_break()

    # ═══ 摘要 ═══
    abst(doc,
        '基于液晶可变延迟器（LCVR）的Stokes偏振测量系统在宽波段工作中面临噪声放大'
        '与温度漂移的双重挑战。现有优化方法通常针对条件数（CN）或泊松噪声传播因子'
        '（BCPN）单目标进行优化，忽略了温度变化对测量精度的显著影响。本文提出以CN、'
        'BCPN和温漂敏感度为三目标的NSGA-II多目标优化框架，将噪声抑制与温漂鲁棒性的'
        '联合优化作为目标。在Haller温漂模型框架下，于350-700nm波段、−10~40°C温度'
        '范围内对6次测量LCVR-Stokes配置进行优化设计。优化所得膝点解实现了CN=2.162、'
        'BCPN=0.471、极差ΔCN=0.089的综合性能。消融仿真分析表明，加入温漂'
        '目标后ΔCN降低75.4%。Monte Carlo Stokes重建误差验证了该配置在宽温域内'
        '具有与无温漂优化方案相当的测量精度。',
        'LCVR；Stokes偏振计；NSGA-II；温漂鲁棒性；多目标优化')
    doc.add_page_break()

    # ═══ 第1章 绪论 ═══
    h1(doc, '第1章  绪论')
    h2(doc, '1.1  研究背景')
    body(doc,
        'Stokes偏振测量技术能够完整获取光束的偏振态信息（Stokes矢量S=[S₀,S₁,S₂,S₃]ᵀ），'
        '在天文观测、遥感探测、生物医学成像和材料表征等领域具有广泛应用[1,15,16,17,21]。'
        '液晶可变延迟器（LCVR）作为一种电控延迟器件，具有无机械运动部件、功耗低、响应快、'
        '孔径大等优势，已成为分时型Stokes偏振测量系统的核心元件[4,5,10]。')
    body(doc,
        '典型的LCVR Stokes偏振测量系统通过改变LCVR的驱动电压来调制入射光的偏振态，'
        '在不同调制状态下测量光强信号，构建测量矩阵并反演Stokes矢量。然而，LCVR的'
        '延迟量不仅依赖于驱动电压和入射波长，还显著受环境温度影响[6,9,18,20]。在'
        '宽波段（350-700 nm）和宽温域（−10~40°C）的实际工作条件下，温度漂移会导致'
        '延迟量偏差，进而引起测量矩阵扰动和Stokes矢量反演精度恶化。')

    h2(doc, '1.2  国内外研究现状')
    body(doc,
        'Stokes偏振测量系统的优化设计是偏振光学的重要研究方向。Sabatke等[1]提出以'
        '条件数（CN）作为偏振计优劣准则。Lara和Paterson[2]将散粒噪声引入优化框架，'
        '提出泊松噪声传播因子（BCPN）。Letnes等[3]采用遗传算法实现宽波段偏振计快速'
        '优化。Peinado等[4]推广到可变延迟器偏振系统并给出实验验证。Wu等[5]针对LCVR'
        '宽波段偏振计进行了系统优化。Chang等[7]提出了LCVR Stokes偏振计的最优目标函数。'
        'Yang等[8]开展了三维仪器偏振分析。Jiang等[10]将LCVR分时偏振成像应用于宽波段'
        '成像。Montes-González等[11]开发了NIR-RGB Mueller矩阵成像偏振计。Chang等[12]'
        '优化了宽波段铁电液晶Stokes偏振计。Bruce和Montes-González[13]分析了实验误差'
        '的影响。Lai等[14]研究了双Stokes-Mueller偏振测量的最优测量矩阵。')
    body(doc,
        '在温漂方面，Niu等[6]研究了LCVR温度漂移抑制方法。Xie等[9]建立了LCVR的非线性'
        '偏振漂移模型与实时补偿方法。Coman和Berg[18]对LCVR偏振特性进行了表征与优化。'
        '然而，上述方法大多仅针对CN或BCPN单一指标进行优化，尚未将温漂敏感度作为独立'
        '优化目标纳入多目标框架。')

    h2(doc, '1.3  研究目的与意义')
    body(doc,
        '针对上述不足，本文提出以条件数（CN）、泊松噪声传播因子（BCPN）和温漂敏感度'
        '为三目标的NSGA-II多目标优化框架[19]，在Haller温漂模型框架下对LCVR Stokes偏振'
        '计进行联合优化设计。据作者检索，目前尚未发现公开报道的针对LCVR Stokes偏振计'
        '考虑CN、BCPN及温漂敏感度的NSGA-II优化研究。本文的主要贡献包括：'
        '（1）将温漂敏感度作为独立优化目标，实现噪声抑制与温度鲁棒性的联合优化；'
        '（2）在6次测量配置下获得了综合性能均衡的膝点解；（3）通过消融仿真分析和温度扫描'
        '分析，验证了多目标优化框架的有效性。')

    # ═══ 第2章 理论模型 ═══
    doc.add_page_break()
    h1(doc, '第2章  理论模型')

    h2(doc, '2.1  LCVR延迟-电压-温度模型')
    body(doc,
        'LCVR的延迟量δ取决于驱动电压V、入射波长λ和环境温度T。液晶的双折射率Δn随'
        '温度的变化遵循Haller经验模型[20]：')
    ctr(doc, 'Δn(T) = Δn₀ · (1 − T/Tc)^β')
    body(doc,
        '其中Δn₀为绝对零度时的外推双折射率，Tc为液晶清亮点温度，β为材料常数。本文'
        '采用Tc=85°C、β=0.162的典型向列相液晶参数[9]。由延迟量与双折射率的关系'
        'δ = 2πd·Δn/λ（d为液晶层厚度），可得温度依赖的延迟量：')
    ctr(doc, 'δ(λ,V,T) = [2πd/λ] · Δn₀ · x(V) · (1 − T/Tc)^β')
    body(doc,
        '其中x(V)为电压依赖的指向矢取向因子。将(1 − T/Tc)^β在参考温度T₀附近进行'
        '一阶泰勒展开，并令δ₀(λ,V)为T₀下的延迟量，可得线性近似表达式：')
    ctr(doc, 'δ(λ,V,T) ≈ δ₀(λ,V) · [1 + α(T − T₀)]')
    body(doc,
        '其中等效温漂系数α = −β/(Tc − T₀)。本文取T₀=25°C，得α = −0.162/60 = '
        '−0.0027 °C⁻¹。根据文献[9]，该线性近似在−10~40°C范围内可满足工程计算要求。'
        '后续分析基于此温漂线性模型，在−10~40°C范围内对6次测量LCVR-Stokes配置进行'
        '优化设计。')

    h2(doc, '2.2  测量矩阵构建')
    body(doc,
        'LCVR双液晶Stokes偏振测量系统由线性偏振器、两个LCVR（LCVR₁和LCVR₂）和探测器'
        '组成。两个LCVR的快轴方位角分别为θ₁和θ₂。液晶可变延迟器的穆勒矩阵为：')
    ctr(doc,
        'M_LCVR(δ,θ) = [1, 0, 0, 0;\n'
        ' 0, cos²2θ+sin²2θ·cosδ, cos2θ·sin2θ·(1−cosδ), −sin2θ·sinδ;\n'
        ' 0, cos2θ·sin2θ·(1−cosδ), sin²2θ+cos²2θ·cosδ, cos2θ·sinδ;\n'
        ' 0, sin2θ·sinδ, −cos2θ·sinδ, cosδ]',
        sz=Pt(10))
    body(doc,
        '6次测量配置下对应6组电压组合(V₁ᵢ,V₂ⱼ)，i=1,2,3, j=1,2。系统穆勒矩阵为'
        'M_sys = M_polarizer · M_LCVR₂(δ₂ⱼ,θ₂) · M_LCVR₁(δ₁ᵢ,θ₁)。第k次测量光强'
        '为M_sys第一行与入射Stokes矢量的内积，由此得到6×4测量矩阵A的第k行表达式：')
    ctr(doc, 'A_k = [1, a₁ᵢⱼ, a₂ᵢⱼ, a₃ᵢⱼ]')
    body(doc, '其中显式三角函数表达式为（c₂ᵢ=cos2θᵢ, s₂ᵢ=sin2θᵢ）：')
    ctr(doc,
        'a₁ = (c₂₂²+s₂₂²·cδ₂ⱼ)(c₂₁²+s₂₁²·cδ₁ᵢ)\n'
        '     + c₂₂·s₂₂·(1−cδ₂ⱼ)·c₂₁·s₂₁·(1−cδ₁ᵢ) − s₂₂·sδ₂ⱼ·s₂₁·sδ₁ᵢ',
        sz=Pt(10))
    ctr(doc,
        'a₂ = (c₂₂²+s₂₂²·cδ₂ⱼ)·c₂₁·s₂₁·(1−cδ₁ᵢ)\n'
        '     + c₂₂·s₂₂·(1−cδ₂ⱼ)·(s₂₁²+c₂₁²·cδ₁ᵢ) + s₂₂·sδ₂ⱼ·c₂₁·sδ₁ᵢ',
        sz=Pt(10))
    ctr(doc,
        'a₃ = (c₂₂²+s₂₂²·cδ₂ⱼ)(−s₂₁·sδ₁ᵢ)\n'
        '     + c₂₂·s₂₂·(1−cδ₂ⱼ)·c₂₁·sδ₁ᵢ − s₂₂·sδ₂ⱼ·cδ₁ᵢ',
        sz=Pt(10))
    body(doc,
        'Stokes矢量由最小二乘法反演。定义Moore-Penrose伪逆矩阵A⁺ = (AᵀA)⁻¹Aᵀ，则Stokes矢量的最小二乘估计为：')
    ctr(doc, 'S = A⁺I    （2）')
    body(doc,
        '其中Stokes矢量S=[S₀,S₁,S₂,S₃]ᵀ，I为6维光强测量向量。式（2）与直接求解正规方程S = (AᵀA)⁻¹AᵀI等价，A⁺各行的2-范数即为后续计算EWV的基础。')

    h2(doc, '2.3  噪声传播指标')
    body(doc,
        '条件数（CN）表征测量矩阵对高斯加性噪声的鲁棒性[1,7]：κ(A)=σ_max/σ_min。'
        '泊松噪声传播因子BCPN[2,22]定义为q_sum=Σᵢ||Qᵢ||₂，其中Qᵢ为噪声传播矩阵的'
        'Frobenius范数。等权方差（EWV）[1]定义为EWV=(1/N)Σᵢ(1/σᵢ²)，其中σᵢ为A⁺'
        '各行的2-范数。上述指标均基于350-700nm波段内36个均匀离散波长点和−10~40°C间'
        '11个温度点（步长5°C）的(λ,T)组合扫描计算取算术平均值。36波长和11温度的采样'
        '密度经预实验验证已满足收敛要求（波长点数增至72时CN变化<1%）。')

    h2(doc, '2.4  温漂敏感度指标')
    body(doc,
        '本文定义两种温漂指标：（1）优化阶段采用的CN标准差σ_T = std_T(mean_λ(κ(A)))，'
        '该指标统计稳健，适合驱动Pareto进化；（2）评估阶段采用的极差ΔCN = '
        'max(κ(A(T))) − min(κ(A(T)))，该指标更直观反映温漂幅度。优化目标函数f₃采用'
        'σ_T，最终性能评价采用ΔCN。仿真表明二者具有强相关性（Pearson r>0.95），因此'
        'σ_T的最小化可以有效约束ΔCN。')

    h2(doc, '2.5  NSGA-II优化算法')
    body(doc,
        '决策变量为7维向量：θ₁,θ₂,V₁₁,V₁₂,V₁₃,V₂₁,V₂₂。三目标函数为：'
        'f₁=mean(κ)、f₂=mean(q_sum)、f₃=std_T(mean_λ(κ))。NSGA-II[19]参数：'
        '种群规模100，进化代数200，交叉概率0.8，变异概率0.2。优化完成后采用L-BFGS-B'
        '对Pareto前沿解进行局部精化，以归一化最小欧氏距离法选取膝点。为评估统计稳定性，'
        '本文重复了5次独立NSGA-II运行，膝点解的CN标准差<0.01，BCPN标准差<0.02。')

    # ═══ 第3章 优化结果 ═══
    doc.add_page_break()
    h1(doc, '第3章  优化结果')

    h2(doc, '3.1  综合性能对比')
    body(doc,
        '将单目标GA基准、NSGA-II多目标优化、消融仿真分析与文献数据汇总对比如表1所示。'
        '其中Chang等[7]的数据提取自原文图表，本文单目标GA用于验证模型一致性。需注意，'
        '由于文献[7]和[5]采用的波段范围、测量次数和优化目标与本文不完全相同，表1中文献'
        '数据仅用于量级参考，而非严格同条件比较。'
        '温度范围−10~40°C来自LCVR典型工作环境规格[5,6]（民用/工业级LCVR通常标称'
        '−10~50°C）。')
    tbl(doc,
        ['方法/配置', 'CN', 'BCPN', 'ΔCN', 'EWV', '说明'],
        [
            ['Chang 2024 (4meas)', '~2.38', '~0.9', '—', '—', '单目标CN, 350-700nm'],
            ['Chang 2024 (6meas)', '~2.03', '~0.5', '—', '—', ''],
            ['Wu 2022 LCVR', '~2.1', '—', '—', '—', '490-700nm波段'],
            ['本文单目标GA (6m)', '2.030', '0.510', '—', '0.321', '模拟Chang方法'],
            ['本文2目标(无温漂)', '2.094', '0.349', '0.362', '0.318', 'NSGA-II CN+BCPN'],
            ['本文3目标(有温漂)', '2.162', '0.471', '0.089', '0.325', 'NSGA-II 三目标'],
        ],
        cap='表1  系统配置综合性能对比'
    )
    body(doc,
        '三目标膝点解的CN=2.162较单目标最优2.03高约6.5%，但BCPN从0.51降至0.471'
        '（改善7.6%），且ΔCN=0.089，这是单目标方法无法获得的。')

    h2(doc, '3.2  Pareto前沿与膝点选择')
    body(doc,
        'NSGA-II优化获得的Pareto前沿（100个非支配解）及其膝点如图1所示。膝点经归一化'
        '最小欧氏距离法确定（索引#35），参数为θ₁=−25.23°, θ₂=−49.18°, '
        'V₁=[2.083, 1.570, 4.805]V, V₂=[1.953, 8.699]V。')
    # 插入图1 Pareto前沿
    pp = '/mnt/c/Users/45064/Desktop/pareto_front_3D_CN.png'
    pa = '/home/user2/research-project/lcvr-optimization/figures/pareto_front_3d.png'
    if os.path.exists(pp):
        doc.add_picture(pp, width=Cm(12)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif os.path.exists(pa):
        doc.add_picture(pa, width=Cm(12)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    figcap(doc, '图1  NSGA-II三维Pareto前沿（★为膝点，f₁=CN-e4, f₂=BCPN, f₃=Drift std）')
    # 插入2D投影图
    pp2 = '/mnt/c/Users/45064/Desktop/pareto_front_2D_CN.png'
    pa2 = '/home/user2/research-project/lcvr-optimization/figures/pareto_front_2d_projections.png'
    if os.path.exists(pp2):
        doc.add_picture(pp2, width=Cm(14)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif os.path.exists(pa2):
        doc.add_picture(pa2, width=Cm(14)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    figcap(doc, '图2  Pareto前沿二维投影 (a) f₁-f₂ (b) f₁-f₃ (c) f₂-f₃')

    h2(doc, '3.3  消融仿真分析')
    tbl(doc,
        ['方法', 'CN', 'BCPN', 'ΔCN'],
        [
            ['2目标（无温漂）', '2.094', '0.349', '0.362'],
            ['3目标（有温漂）', '2.162', '0.471', '0.089'],
            ['变化', '↑3.2%', '↑35%', '↓75.4%'],
        ],
        cap='表2  消融仿真分析对比'
    )
    body(doc,
        '消融仿真分析表明加入温漂目标后ΔCN降低75.4%，CN仅增3.2%。BCPN上升35%的物理本质'
        '在于：温漂优化迫使六个调制状态在温度变化下保持矩阵形状稳定，各状态在庞加莱球'
        '上的分布偏离了散粒噪声最优的正交配置，导致泊松噪声传播因子增大。图4给出了'
        '三种优化方案在庞加莱球上调制态分布的对比，3目标解的调制态分布更为紧凑，体现'
        '了温度稳定性的代价。这一现象反映了温漂鲁棒性与泊松噪声抑制之间存在明显的权衡关系。')
    # 插入Poincaré球图
    pp3 = '/mnt/c/Users/45064/Desktop/poincare_sphere.png'
    pa3 = '/home/user2/research-project/lcvr-optimization/figures/poincare_sphere.png'
    if os.path.exists(pp3):
        doc.add_picture(pp3, width=Cm(14)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif os.path.exists(pa3):
        doc.add_picture(pa3, width=Cm(14)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    figcap(doc, '图4  庞加莱球上调制态分布 (a) GA 6meas (b) 2目标 (c) 3目标（★为质心）')

    h2(doc, '3.4  Monte Carlo Stokes误差验证')
    body(doc,
        '为直接评估温漂优化对Stokes重建精度的影响，设计Monte Carlo仿真实验。'
        '随机生成500组覆盖庞加莱球均匀分布的Stokes向量S=[1,S₁,S₂,S₃]，对每组'
        '入射光在−10~40°C范围内11个温度点和36个波长点下计算含高斯噪声和散粒噪声'
        '的测量强度信号，采用各优化配置的伪逆矩阵重建Stokes矢量。噪声水平设定为'
        '信号均值的1%（高斯）和1000光子计数（泊松近似）。全温域平均结果如表3和图5'
        '所示，其中RMSE为500样本×36波长×11温度的全面统计。')
    # 插入MC RMSE图
    mc1 = '/mnt/c/Users/45064/Desktop/mc_stokes_rmse.png'
    mc1b = '/home/user2/research-project/lcvr-optimization/figures/monte_carlo_rmse.png'
    if os.path.exists(mc1):
        doc.add_picture(mc1, width=Cm(14)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif os.path.exists(mc1b):
        doc.add_picture(mc1b, width=Cm(14)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    figcap(doc, '图5  不同温度下Stokes重建RMSE (a) S₁ (b) S₂ (c) S₃ (d) 总RMSE')

    tbl(doc,
        ['方法', 'CN(25°C)', 'BCPN(25°C)', 'ΔCN', 'Total RMSE(25°C)', 'Total RMSE(全温域)'],
        [
            ['GA 6meas¹', '2.010', '0.533', '0.099', '0.01415', '0.01435'],
            ['2目标(无温漂)', '2.094', '0.349', '0.272', '0.01441', '0.01476'],
            ['3目标(有温漂)', '2.162', '0.471', '0.063', '0.01436', '0.01451'],
        ],
        cap='表3  Monte Carlo Stokes重建误差对比（500样本/温度点×36波长，基于非线性Haller模型评估）\n'
             '¹单目标CN优化的GA基线'
    )
    body(doc,
        '表3和图5的结果表明：在36波长×11温度的全面评估中，三目标优化解在25°C时'
        'Total RMSE=0.01436，与2目标解（0.01441）和GA基线（0.01415）基本持平；'
        '在全温域（−10~40°C）平均Total RMSE中，3目标解（0.01451）优于2目标解'
        '（0.01476），与GA基线（0.01435）相当。值得注意的是，上述Monte Carlo评估'
        '基于完整非线性Haller温漂模型（式(1)）而非优化阶段采用的线性近似，这验证了'
        '线性近似优化的结果在实际非线性温漂环境中仍然有效。从温漂角度（ΔCN）看，'
        '单目标CN优化的GA基线ΔCN仅为0.099，说明CN优化本身已具有部分温漂抑制效果；'
        '2目标（CN+BCPN）的ΔCN恶化至0.272，这是因为BCPN优化迫使调制态向正交'
        '配置靠拢，却牺牲了温度稳定性；3目标重新将ΔCN降至0.063（较GA降低36.4%，'
        '较2目标降低76.8%），验证了温漂目标约束的有效性。为进一步评估统计显著性，'
        '在550nm单波长下以5000组独立Stokes样本进行了补充分析（采用双侧Welch t-test，'
        '不等方差t检验）。需注意，t-test仅覆盖550nm单波长点，而表3为36波长的全波段'
        '均值。结果表明：在−10°C极端温下，3目标解RMSE=0.01229，2目标解RMSE=0.01233，'
        '差异无统计学意义（p=0.66）；在25°C和40°C下，2目标解RMSE略优（p<0.001），'
        '差异幅度约3-4%。这说明在550nm单波长下3目标优化的代价是在室温下牺牲了少量'
        '精度以换取温度稳定性。全温域平均RMSE中3目标解优于2目标解（0.01451 vs '
        '0.01476），这是因为全温域多波长均值受低温区域（−10~0°C）的CN恶化贡献更大'
        '——在低温区2目标解的CN从2.094上升至2.325（Δ=0.231），而3目标解仅上升至'
        '2.222（Δ=0.060），因此全波段均值反而改善。ΔCN的改善对于保证仪器在宽温域内'
        '稳定运行具有重要的工程意义。')

    # ═══ 第4章 分析 ═══
    doc.add_page_break()
    h1(doc, '第4章  分析与讨论')

    h2(doc, '4.1  温度影响分析')
    body(doc,
        '对2目标和3目标膝点解在−10~40°C范围内进行逐温度评估，结果如图3所示。')
    # 插入温度扫描图
    f4 = '/mnt/c/Users/45064/Desktop/论文/fig4_temperature_scan.png'
    f4b = '/home/user2/research-project/lcvr-optimization/figures/temperature_scan_comparison.png'
    if os.path.exists(f4):
        doc.add_picture(f4, width=Cm(14)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif os.path.exists(f4b):
        doc.add_picture(f4b, width=Cm(14)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    figcap(doc, '图3  不同温度下CN和BCPN变化 (a) CN vs 温度 (b) BCPN vs 温度')
    body(doc,
        '图3(a)显示2目标解在低温段（−10°C）CN恶化至2.417，高温段恢复至2.055，'
        'ΔCN=0.362。这是因为在远离T₀的低温区，(T−T₀)的负值较大，一阶线性近似'
        'α(T−T₀)的绝对值大，未受温漂约束的电压配置偏离最优调制区间，测量矩阵正交性'
        '退化。3目标解全温域CN稳定在2.159~2.248（ΔCN=0.089），因为温漂目标约束'
        '优化过程找到了对温度不敏感的电压-延迟组合区域。图3(b)中3目标BCPN曲线更平坦'
        '（0.459~0.511），虽然均值略高但温域稳定性显著更优。温度扫描数据（−10~40°C，'
        '步长5°C）：2目标CN [2.417, 2.343, 2.281, 2.227, 2.183, 2.147, 2.117, '
        '2.094, 2.075, 2.063, 2.055]；3目标CN [2.248, 2.227, 2.211, 2.197, 2.185, '
        '2.175, 2.168, 2.162, 2.160, 2.159, 2.161]。')

    h2(doc, '4.2  创新点与讨论')
    body(doc,
        '本文的主要创新点包括：（1）将CN、BCPN和温漂敏感度作为三目标纳入NSGA-II框架，'
        '实现了噪声抑制与温度鲁棒性的联合优化；（2）在Haller温漂理论框架下获得了宽波段、'
        '宽温域内性能均衡的6次测量最优配置；（3）系统性的消融仿真分析表明温漂鲁棒性与泊松'
        '噪声抑制之间存在明显权衡关系。')
    body(doc,
        '本框架具有良好的通用性：对于不同型号或Haller参数的LCVR元件，只需将实测参数'
        '输入本框架即可快速定制特定器件的宽温域最优配置。当前研究的局限在于：温漂指标'
        'ΔCN是测量矩阵条件数的稳定性度量，而非Stokes反演误差的直接度量（后者需引入'
        '具体的噪声模型和入射光偏振态）。未来工作将结合实验标定数据和Poincaré球调制态'
        '分布分析开展实验验证，并探索自适应温漂补偿方法。')

    # ═══ 第5章 结论 ═══
    doc.add_page_break()
    h1(doc, '第5章  结论')
    body(doc,
        '本文针对LCVR Stokes偏振测量系统的噪声与温漂耦合问题，提出了CN+BCPN+温漂敏感度'
        '三目标NSGA-II优化框架。主要结论如下：')
    body(doc,
        '（1）从Haller双折射率模型出发，经一阶泰勒展开建立了延迟量的线性温漂近似模型，'
        '在350-700nm、−10~40°C范围内对6次测量配置进行了优化。膝点解参数θ₁=−25.23°, '
        'θ₂=−49.18°, V₁=[2.083, 1.570, 4.805]V, V₂=[1.953, 8.699]V，实现了'
        'CN=2.162、BCPN=0.471、ΔCN=0.089的综合性能。')
    body(doc,
        '（2）消融仿真分析证明加入温漂目标后ΔCN降低75.4%（0.362→0.089），CN仅增3.2%。'
        'BCPN的35%增幅反映了温漂鲁棒性与泊松噪声抑制之间存在明显的权衡关系。Monte Carlo'
        'Stokes重建误差验证了3目标解在全温域内具有与2目标解相当的测量精度。')
    body(doc,
        '（3）温度扫描分析表明三目标解全温域CN波动仅0.089，远优于2目标的0.362。')
    body(doc,
        '（4）与单目标GA及文献结果的对比展现了多目标框架在平衡多种噪声源方面的潜在优势，'
        '为宽波段、宽温域LCVR偏振测量系统的工程设计提供了新思路。')

    # ═══ 参考文献 ═══
    doc.add_page_break()
    refs_list = [
        '[1] Sabatke D S, Descour M R, Dereniak E L, et al. Figures of merit for complete Stokes polarimeter optimization[J]. Applied Optics, 2000, 39(32): 5935-5939.',
        '[2] Lara D, Paterson C. Stokes polarimeter optimization in shot and Gaussian noise[J]. Optics Express, 2009, 17(23): 20672-20687.',
        '[3] Letnes P A, Frenning G, Ross M, et al. Fast and optimal broad-band Stokes/Mueller polarimeter design by genetic algorithm[J]. Optics Express, 2010, 18(22): 23095-23103.',
        '[4] Peinado A, Lizana A, Campos J, et al. Analysis, optimization and implementation of a variable retardance based polarimeter[C]. SPIE, 2010, 7782: 77820A.',
        '[5] Wu C, Zhang J, Li Y, et al. Optimal design for a broadband Stokes polarimeter of liquid crystal variable retarders[J]. Applied Optics, 2022, 61(36): 10685-10693.',
        '[6] Niu Y, Duan L, Zhang J, et al. Suppression of ambient temperature-caused drift in a laser power stabilization system with a liquid crystal variable retarder in atomic gyroscopes[J]. Review of Scientific Instruments, 2022, 93(4): 043002.',
        '[7] Chang L, Li J, Zhang Y, et al. Optimization analysis of a Stokes polarimeter for broadband liquid crystal variable retarders under the optimal objective function[J]. Journal of the Optical Society of America A, 2024, 41(4): 606-615.',
        '[8] Yang J, Zhang X, Hou J, et al. Three-dimensional instrument polarization analysis and optimization of liquid-crystal-based Stokes polarimeter[J]. Applied Optics, 2023, 62(36): 9612-9621.',
        '[9] Xie H, Zhou X, Cao Z, et al. Nonlinear polarization drift modeling and real-time compensation in LCVR-based laser power stabilization systems[J]. Applied Optics, 2025, 64(29): 8633-8642.',
        '[10] Jiang W, Ren L, Liang J. LCVR-based broadband time-division polarimetric imaging technology in visible spectral range[J]. Optics and Lasers in Engineering, 2026, 201: 109581.',
        '[11] Montes-González I, Bisbal J, Li D, et al. Optimization and development of a compact NIR-RGB Mueller matrix imaging polarimeter: application to in situ skin measurements[J]. Advanced Photonics Nexus, 2026, 5(2): 026009.',
        '[12] Chang L, Liu J, Yin Y, et al. Optimization of broadband ferroelectric liquid crystal Stokes polarimeters: evaluation and selection[C]. AOPC 2025: Optical Sensing, Imaging, Communications, Display, and Biomedical Optics, SPIE, 2025, 13958: 23.',
        '[13] Bruce N C, Montes González I. The effect of experimental errors on optimized polarimeters[C]. Polarization Science and Remote Sensing XI, SPIE, 2023, 12690: 24.',
        '[14] Lai C, Li Y, Wu J, et al. The optimal measurement matrix with minimal condition number for double Stokes-Mueller polarimetry[C]. AOPC 2022: Biomedical Optics, SPIE, 2023, 12560: 13.',
        '[15] Azzam R M A. Stokes-vector and Mueller-matrix polarimetry[J]. Journal of the Optical Society of America A, 2016, 33(7): 1396-1408.',
        '[16] Tyo J S, Goldstein D H, Chenault D B, et al. Review of passive imaging polarimetry for remote sensing applications[J]. Applied Optics, 2006, 45(22): 5453-5469.',
        '[17] Goldstein D H. Polarized Light[M]. 3rd ed. Boca Raton: CRC Press, 2011.',
        '[18] Coman A, Berg M. Liquid crystal variable retarders for polarimetry: characterization and optimization[J]. Optics & Laser Technology, 2020, 132: 106462.',
        '[19] Deb K, Pratap A, Agarwal S, et al. A fast and elitist multiobjective genetic algorithm: NSGA-II[J]. IEEE Transactions on Evolutionary Computation, 2002, 6(2): 182-197.',
        '[20] Haller I. Thermodynamic and static properties of liquid crystals[J]. Solid State Physics, 1975, 14: 213-248.',
        '[21] Bueno J M. Polarimetry in the human eye: a review[J]. Journal of Optics A, 2001, 3(3): R1-R10.',
        '[22] Goudail F, Bénière A. Estimation precision of the degree of linear polarization and of the angle of polarization in the presence of additive Gaussian and Poisson noise[J]. Optics Express, 2009, 17(17): 14602-14614.',
    ]
    refs(doc, refs_list)
    doc.add_page_break()

def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(3); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3); s.right_margin = Cm(2.5)
    sty = doc.styles['Normal']
    sty.font.name = 'Times New Roman'; sty.font.size = Pt(12)
    sty.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    sty.paragraph_format.line_spacing = 1.5
    gen(doc)
    out = os.path.expanduser('~/research-project/lcvr-optimization/paper/manuscript_zh_v4.docx')
    doc.save(out)
    print(f'OK: {out}')
    import shutil
    for dst in ['/mnt/c/Users/45064/Desktop/LCVR_Paper_V5.1.docx',
                '/mnt/c/Users/45064/Desktop/论文/LCVR_Paper_V5.1.docx']:
        try:
            shutil.copy2(out, dst)
            print(f'Copied: {dst}')
        except:
            print(f'Skip: {dst}')

if __name__ == '__main__':
    main()
