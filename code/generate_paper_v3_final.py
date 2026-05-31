#!/usr/bin/env python3
"""生成V3论文 — 规范公式排版 + GB/T格式 + EWV对比 + 新文献"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, json

doc = Document()
fig_dir = "/tmp/paper_figs"

# ====== 页面设置 (GB/T标准) ======
for section in doc.sections:
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

def set_font(run, cn='宋体', en='Times New Roman', size=Pt(12), bold=False):
    run.font.name = en; run.font.size = size; run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), cn)

def add_h1(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_before = Pt(12); pf.space_after = Pt(6); pf.line_spacing = 1.5
    r = p.add_run(text); set_font(r, '黑体', '黑体', Pt(16), True)

def add_h2(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format; pf.space_before = Pt(8); pf.space_after = Pt(4); pf.line_spacing = 1.5
    r = p.add_run(text); set_font(r, '黑体', '黑体', Pt(14), True)

def add_body(doc, text, indent=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    if indent: pf.first_line_indent = Pt(24)
    pf.line_spacing = 1.5
    r = p.add_run(text); set_font(r)

def add_formula(doc, text, fid=""):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_before = Pt(6); pf.space_after = Pt(6)
    r = p.add_run(text); set_font(r, '宋体', 'Times New Roman', Pt(11))
    r.font.italic = True

def add_formula_inline(doc, text):
    r = doc.add_paragraph().add_run(text)
    set_font(r, '宋体', 'Times New Roman', Pt(11))
    r.font.italic = True
    return r

def make_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = ''; p = t.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); set_font(r, '黑体', 'Times New Roman', Pt(9), True)
    for rd in rows:
        rc = t.add_row().cells
        for i, v in enumerate(rd):
            rc[i].text = ''; p = rc[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(str(v)); set_font(r, '宋体', 'Times New Roman', Pt(9))
    return t

# ====== 标题 ======
for _ in range(2): doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('双LCVR Stokes偏振测量系统的多目标优化设计'); set_font(r, '黑体', '黑体', Pt(20), True)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('——兼顾高斯噪声、泊松噪声与温度漂移的测量配置优化')
set_font(r2, '楷体', 'Times New Roman', Pt(14)); r2.font.color.rgb = RGBColor(80, 80, 80)
doc.add_paragraph()

# ====== 摘要 ======
add_h1(doc, '摘要')
add_body(doc, '针对双液晶可变延迟器(LCVR) Stokes偏振测量系统，本文提出了一种综合考虑高斯噪声、泊松噪声和温度漂移的测量配置优化方法。系统工作在350–700 nm波段，以六次过采样(2×3)方案优化7个可调参数（快轴角度2个、驱动电压5个）。采用条件数κ₂的e4准则、BCPN（Bias-Conditioned Poisson Noise）指标和温度间条件数标准差作为三个优化目标，使用NSGA-II算法求解Pareto前沿。与随机初始配置相比，优化后的膝点配置将Stokes向量重建误差从0.0245降至0.0091，快轴角度在±3°误差内条件数变化小于5%。')
add_body(doc, '需要指出，本文的优化结果基于器件参数模型（LCVR延迟量温漂系数固定为−0.0027/°C），实际器件的非线性和制造误差可能带来额外偏差。')
p = doc.add_paragraph(); pf = p.paragraph_format; pf.first_line_indent = Pt(24); pf.line_spacing = 1.5
r = p.add_run('关键词：'); set_font(r, '黑体', 'Times New Roman', Pt(12), True)
r2 = p.add_run('Stokes偏振测量；液晶可变延迟器；多目标优化；噪声传播；温漂'); set_font(r2)

# ====== 1. 引言 ======
add_h1(doc, '1 引言')
add_h2(doc, '1.1 研究背景')
add_body(doc, '全Stokes偏振测量在遥感探测、生物医学成像和材料表征等领域有广泛应用[1-3]。基于液晶可变延迟器（Liquid Crystal Variable Retarder, LCVR）的偏振计具有无机械运动部件、响应速度快、结构紧凑等优点[4,5]，但测量精度受高斯探测器噪声、泊松光子噪声、LCVR延迟量的温度漂移[6,7]以及快轴装配误差[8]等多种因素影响。')

add_h2(doc, '1.2 相关工作')
add_body(doc, '偏振测量系统的优化设计是偏振光学领域的核心问题之一。Sabatke等人[9]于2000年系统提出了完整Stokes偏振计的评价指标，包括条件数（Condition Number, CN）和等权方差（Equally Weighted Variance, EWV），并给出了四种测量配置下CN的理论下限√3。Lara和Paterson[10]于2009年将优化扩展至同时存在高斯噪声和散粒噪声的场景。Goudail[11]于2013年提出了BCPN指标用于量化泊松噪声下的偏振参数重建偏差。')
add_body(doc, 'Letnes等人[12]在2010年使用遗传算法在430–2000 nm波段对铁电液晶（FLC）和波片组合的偏振计进行了优化，以CN为目标获得了κ₂≈2.2–2.5的优化配置。Peinado等人[13]在同一年对基于双LCVR的Stokes偏振计进行了系统的CN和EWV对比优化，并在单一波长（633 nm）下进行了实验验证。')
add_body(doc, 'Wu等人[14]在2022年针对双LCVR Stokes偏振计在490–700 nm波段提出了一种三步优化方法。Chang、Li等人[15]在2024年将上述工作扩展至350–700 nm波段，解析推导了最优目标函数，并采用遗传算法同时针对高斯噪声和泊松噪声进行了优化分析，讨论了测量次数与整体性能的关系。Jiang等人[16]在2026年报道了基于LCVR的可见光宽带分时偏振成像技术。')
add_body(doc, '在LCVR温度效应方面，Niu等人[6]在2022年实验测量了LCVR在原子陀螺仪激光稳功率系统中的温度漂移。Xie等人[7]在2025年进一步建立了LCVR的非线性温漂模型并实现了实时补偿。Bruce和Montes-González[8]在2023年证明单纯以CN最小化为目标的优化配置不一定在实验误差存在时具有最佳稳定性。Montes-González和Lizana等人[17]在2026年报道了便携式LCVR Mueller矩阵成像偏振计。Chang等人[18]在2025年进一步研究了宽带FLC Stokes偏振计的优化。')
add_body(doc, '纵观上述研究，现有工作普遍采用单目标优化策略（CN或BCPN单一指标），或者对多个噪声场景分步优化。尚无研究将高斯噪声、泊松噪声和温度漂移纳入统一的Pareto多目标优化框架。这正是本文的研究动机。')

add_h2(doc, '1.3 本文工作')
add_body(doc, '本文将高斯噪声、泊松噪声和温度漂移三个指标纳入统一的优化框架，系统比较了四次测量(2×2)和六次测量(2×3)两种方案。值得注意的是，本文采用的优化指标均为已有工作的组合应用，主要贡献在于建立了多目标权衡的评价方法和系统性的鲁棒性验证体系。与现有文献的系统对比见第4.1节表5。')

# ====== 2. 原理与方法 ======
add_h1(doc, '2 原理与方法')
add_h2(doc, '2.1 双LCVR Stokes偏振计的测量模型')
add_body(doc, '双LCVR Stokes偏振计由两个相位延迟可调的LCVR（快轴方位角分别为θ₁和θ₂）后接一个0°方向固定线偏振片组成。入射Stokes向量S = [I, Q, U, V]ᵀ（以光强I归一化）经过LCVR和偏振片后，探测器接收到的光强为：')
add_formula(doc, 'I_meas = A · S + n')
add_body(doc, '其中A为n×4测量矩阵，n为测量次数（4或6）。Stokes向量的估计通过伪逆求解：')
add_formula(doc, 'Ŝ = A⁺ · I_meas，其中A⁺ = (AᵀA)⁻¹Aᵀ')
add_body(doc, '测量矩阵A的行向量由LCVR的Mueller矩阵元素构成。以行向量a_k = [1, m₁₂, m₁₃, m₁₄]为例，其元素表达式为：')
add_formula(doc, 'm₁₂ = (c₂²+s₂²C₂)(c₁²+s₁²C₁) + c₂s₂(1−C₂)·c₁s₁(1−C₁) − s₂S₂·s₁S₁')
add_formula(doc, 'm₁₃ = (c₂²+s₂²C₂)·c₁s₁(1−C₁) + c₂s₂(1−C₂)(s₁²+c₁²C₁) + s₂S₂·c₁S₁')
add_formula(doc, 'm₁₄ = (c₂²+s₂²C₂)(−s₁S₁) + c₂s₂(1−C₂)·c₁S₁ − s₂S₂·C₁')
add_body(doc, '其中c_i = cos(2θ_i), s_i = sin(2θ_i), C_i = cos(δ_i), S_i = sin(δ_i)，δ_i为LCVR的相位延迟量。快轴角度θ_i的优化范围设为[−85°, 85°]（预留5°装配公差边界）。')

add_h2(doc, '2.2 LCVR延迟量模型')
add_body(doc, 'LCVR延迟量δ由双折射色散、电压响应和温度漂移三部分构成：')
add_formula(doc, 'δ(λ,V,T) = 360° × R(λ,V) × [1 + α(T − T₀)] / λ')
add_body(doc, '其中双折射光程差R(λ,V)结合了材料色散和液晶分子取向：')
add_formula(doc, 'R(λ,V) = (c₀ + c₁/λ² + c₂/λ⁴) × [A₂ + (A₁−A₂)/(1+(V/X₀)ᴾ)]')
add_body(doc, 'c₀=4.9538×10², c₁=1.6710×10⁸, c₂=−2.1204×10¹³（λ单位为nm）为双折射色散系数，A₁=1.06572, A₂=0.06539, X₀=1.46116 V, P=3.36309为电压响应曲线参数（取自器件数据手册）。温度系数α=−0.0027/°C（T₀=25°C）来自文献[6]的实测结果，该值在20–50°C范围内近似线性。Xie等人[7]指出，在更宽温度范围或高精度应用中需改用非线性多项式模型。')

add_h2(doc, '2.3 优化目标的形式化定义')
add_body(doc, '目标1（高斯噪声传播）：条件数κ₂的e4偏差。条件数CN = κ₂(A) = ‖A‖₂·‖A⁺‖₂是表征高斯噪声传播放大的经典指标[9]。全Stokes偏振测量（m=4）的CN理论下限为√3≈1.732[9]。Sabatke等人[9]同时提出了等权方差EWV = (1/n) Σ_i (1/s_i²)作为替代评价标准，其中s_i为A的奇异值。对于完整Stokes系统，最小化CN与最大化EWV在优化方向上是等价的[9,10]。本文采用CN的e4准则：')
add_formula(doc, 'f₁ = (1/N_λ) · Σ_λ [1/√3 − 1/κ₂(λ)]⁴')
add_body(doc, '该准则对大偏差施加四次方惩罚，避免仅优化平均条件数而忽视个别波长的劣化。')

add_body(doc, '目标2（泊松噪声传播）：BCPN。定义Q_ij = Σ_k[(A⁺)_ik²]·A_kj（i,j=0,1,2,3），则q_k = (Q_k1, Q_k2, Q_k3)为从Q第k行去掉第一列后的三维向量：')
add_formula(doc, 'f₂ = (1/N_λ) · Σ_λ Σ_{k=0}^{3} ‖q_k(λ)‖₂')
add_body(doc, 'BCPN的物理含义是在光子计数服从泊松分布的假设下，各Stokes参数重建值的偏差系数，其计算方式参见Goudail[11]和Lara等人[10]。')

add_body(doc, '目标3（温度稳定性）：条件数的温度间标准差。温度漂移是LCVR系统长期稳定性的主要制约因素[6,7]。定义：')
add_formula(doc, 'f₃ = (1/N_λ) · Σ_λ σ_T[κ₂(λ,T)],  T ∈ {20°C, 35°C, 50°C}')
add_body(doc, 'σ_T为样本标准差（除以N−1）。f₃越小表示系统在温度变化下性能越稳定。温度点覆盖了实验室空调下限（20°C）到工业设备上限（50°C）。')

add_body(doc, '多目标问题的表述和目标归一化：三个目标量纲和数值范围不同，直接加权无意义。本文采用先使用NSGA-II求解Pareto前沿，然后对Pareto解集内各目标做最小-最大归一化：')
add_formula(doc, "f_i' = (f_i − f_i_min) / (f_i_max − f_i_min)")
add_body(doc, "再取欧氏距离‖[f₁', f₂', f₃']‖₂最小的点为膝点。归一化仅在Pareto解集内部进行，不影响算法寻优过程。")

add_h2(doc, '2.4 NSGA-II算法配置与收敛性')
add_body(doc, '采用NSGA-II[19]求解三目标优化问题。决策变量7个：θ₁,θ₂∈[−85°,85°]，V₁₁,V₁₂,V₁₃,V₂₁,V₂₂∈[1.5 V, 10.0 V]。算法参数：种群规模100，迭代200代，模拟二进制交叉（SBX, η_c=15, p_c=0.9），多项式变异（η_m=20, p_m=1/7）。')
add_body(doc, '对照实验：使用均匀随机采样（各100组随机θ和V组合）作为baseline。结果显示4次测量和6次测量的随机配置中，分别有98%和95%的样本条件数超过100（矩阵接近奇异），前10%最优随机样本的平均条件数分别为28.9（4次）和13.4（6次），远高于优化后的值（κ₂≈2.0–2.3）。这说明可用配置只占很小比例，智能优化算法是必要的。')

if os.path.exists(f"{fig_dir}/fig_convergence.png"):
    doc.add_picture(f"{fig_dir}/fig_convergence.png", width=Inches(5.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run('图1  NSGA-II收敛性：（a）f₁最小值 （b）Pareto前沿规模'); set_font(r, '宋体', 'Times New Roman', Pt(10.5), True)

# ====== 3. 实验结果 ======
add_h1(doc, '3 实验结果')
add_h2(doc, '3.1 单目标优化基线')
add_body(doc, '使用差分进化算法对三个目标分别进行单目标优化，作为性能上界/下界的参考。结果如表1所示。')

make_table(doc, ['配置', 'κ₂(均值±1σ)', 'BCPN', '温漂f₃', 'EWV', '说明'],
    [['4meas CN最优', '2.346±0.494', '1.904', '0.145', '0.783', '仅优化κ₂'],
     ['6meas CN最优', '2.026±0.153', '0.509', '0.053', '0.305', '6变量→7变量'],
     ['6meas BCPN最优', '2.190±0.338', '0.350', '0.122', '0.325', 'BCPN最低但温漂大'],
     ['6meas 膝点★', '2.245±0.065', '0.472', '0.035', '0.318', '多目标膝点']])
cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cp.add_run('表1  单目标优化与膝点对比（含EWV）'); set_font(r, '宋体', 'Times New Roman', Pt(10.5), True)
add_body(doc, '说明：EWV值为全波段(350–700 nm)及多温度(20°C, 35°C, 50°C)平均，数值越小表示高斯噪声传播越小。6次测量BCPN最优解的温漂（0.122）劣于CN最优解（0.053），各指标间的冲突关系正是进行多目标优化的动机。与Letnes等人[12]在430–2000 nm波段得到的最优CN（~2.2–2.5）相比，本文6次测量CN最优（2.026±0.153）略优。')

add_h2(doc, '3.2 多目标优化结果')
add_body(doc, 'NSGA-II经过200代优化得到148个Pareto非支配解（过滤了2个矩阵接近奇异的异常点）。图2展示了Pareto前沿的三维分布。')
if os.path.exists(f"{fig_dir}/fig_pareto.png"):
    doc.add_picture(f"{fig_dir}/fig_pareto.png", width=Inches(5.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run('图2  Pareto前沿及膝点位置（★）'); set_font(r, '宋体', 'Times New Roman', Pt(10.5), True)

add_body(doc, '膝点配置参数见表2。')
make_table(doc, ['参数', '数值'],
    [['θ₁, θ₂', '−65.74°, −41.29°'],
     ['V₁₁, V₁₂, V₁₃', '8.44 V, 1.59 V, 2.06 V'],
     ['V₂₁, V₂₂', '9.72 V, 1.95 V'],
     ['κ₂, BCPN, f₃, EWV', '2.245, 0.472, 0.035, 0.318'],
     ['归一化距离', '0.41（三目标归一化欧氏距离）']])
cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cp.add_run('表2  膝点配置'); set_font(r, '宋体', 'Times New Roman', Pt(10.5), True)

add_h2(doc, '3.3 Stokes向量重建精度')
add_body(doc, '仿真设置：在Poincaré球上均匀采样82个Stokes向量，测量噪声包括高斯噪声（σ=0.001）和泊松噪声，采用最小二乘伪逆重建。结果如表3所示。')
make_table(doc, ['配置', 'Stokes误差', 'DoLP误差', '|dQ|均值', '|dV|均值'],
    [['4meas 随机', '0.0245', '0.0224', '0.0172', '0.0116'],
     ['4meas CN最优', '0.0116', '0.0085', '0.0078', '0.0075'],
     ['6meas CN最优', '0.0089', '0.0065', '0.0060', '0.0058'],
     ['6meas 膝点', '0.0091', '0.0065', '0.0060', '0.0060']])
cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cp.add_run('表3  Stokes重建误差对比'); set_font(r, '宋体', 'Times New Roman', Pt(10.5), True)
add_body(doc, '计算口径说明：本文中的"降低百分比"均以"4次随机配置"为基准。如使用"4次CN最优"为基准，膝点的改善幅度为22% = (0.0116−0.0091)/0.0116×100%。')

if os.path.exists(f"{fig_dir}/fig_noise_scan.png"):
    doc.add_picture(f"{fig_dir}/fig_noise_scan.png", width=Inches(4.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run('图3  不同噪声水平下Stokes重建误差'); set_font(r, '宋体', 'Times New Roman', Pt(10.5), True)

add_h2(doc, '3.4 快轴角度误差鲁棒性')
add_body(doc, '蒙特卡洛仿真：每个误差水平进行50次独立采样，在标称角度上叠加均匀分布随机误差U(−ε, +ε)。结果见表4。')
make_table(doc, ['误差', '4meas κ₂', '膝点 κ₂', '4meas BCPN', '膝点 BCPN'],
    [['0°', '2.293±0.339', '2.263±0.068', '1.786', '0.431'],
     ['±1°', '2.464±0.462', '2.232±0.075', '1.974', '0.473'],
     ['±3°', '2.532±0.474', '2.283±0.138', '1.986', '0.468'],
     ['±5°', '2.684±0.533', '2.319±0.202', '2.108', '0.561']])
cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cp.add_run('表4  快轴角度误差鲁棒性对比'); set_font(r, '宋体', 'Times New Roman', Pt(10.5), True)
add_body(doc, '膝点配置在±3°误差内κ₂变化小于5%，在±5°时变化仅2.5%。四次测量基线在±5°时变化幅度为17.1%。Bruce等人[8]指出，单纯以CN最小化为目标的配置往往对实验误差更敏感，本文多目标优化的膝点配置在角度误差下的稳定性部分验证了这一观点。')

if os.path.exists(f"{fig_dir}/fig_angle_error.png"):
    doc.add_picture(f"{fig_dir}/fig_angle_error.png", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run('图4  快轴角度误差鲁棒性'); set_font(r, '宋体', 'Times New Roman', Pt(10.5), True)

add_h2(doc, '3.5 宽波段与温度稳定性')
add_body(doc, '图5展示了膝点配置在350–700 nm范围内条件数随波长的变化曲线以及15–55°C连续温度扫描结果。在全部温度范围内，膝点条件数仅从2.288变化到2.234，波动幅度为2.4%，BCPN从0.439变化到0.546。')
if os.path.exists(f"{fig_dir}/fig_temp_scan.png"):
    doc.add_picture(f"{fig_dir}/fig_temp_scan.png", width=Inches(5.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run('图5  膝点15–55°C连续温度扫描'); set_font(r, '宋体', 'Times New Roman', Pt(10.5), True)

# ====== 4. 讨论 ======
add_h1(doc, '4 讨论')
add_h2(doc, '4.1 与已有工作的系统对比')
add_body(doc, '表5从多个维度将本文工作与已有文献进行系统对比，新增EWV指标作为高斯噪声传播的补充评价标准。')
make_table(doc, ['文献', '系统', '多目标', '温漂', 'EWV', '角度鲁棒性', '实验'],
    [['Sabatke 2000[9]', '波片', '✗', '✗', '✓', '✗', '✗'],
     ['Lara 2009[10]', '可变延迟器', '✗', '✗', '✗', '✗', '✓'],
     ['Letnes 2010[12]', 'FLC', '✗', '✗', '✗', '✗', '✓'],
     ['Peinado 2010[13]', '双LCVR', '✗', '✗', '✓', '△', '✓'],
     ['Wu 2022[14]', '双LCVR', '✗', '✗', '✗', '✗', '✗'],
     ['Chang 2024[15]', '双LCVR', '✗', '✗', '✗', '✗', '✗'],
     ['本文', '双LCVR', '✓', '✓', '✓', '✓', '✗']])
cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cp.add_run('表5  与已有工作的系统对比'); set_font(r, '宋体', 'Times New Roman', Pt(10.5), True)

add_h2(doc, '4.2 结果对比')
add_body(doc, '与Chang等人[15]的对比：两者波段完全相同（350–700 nm），系统相同（双LCVR）。本文采用单目标GA模拟了Chang的方法，结果为4测量CN=2.38、6测量CN=2.03，与Chang报告的量级（CN≈2.1–2.5）一致，验证了物理模型的可靠性。本文NSGA-II膝点（CN=2.245, BCPN=0.472, 温漂=0.035）在CN上略差于单目标最优，但在BCPN和温漂上取得了更好的平衡。')
add_body(doc, '与Wu等人[14]的结果相比（CN≈2.1–2.3, 490–700 nm），本文膝点CN=2.245处于相同水平。与Letnes等人[12]相比（κ₂≈2.2–2.5, 430–2000 nm），本文在较窄波段取得了相近的CN水平，同时额外考虑了泊松噪声和温漂。')
add_body(doc, 'EWV对比：本文膝点EWV=0.318，与单目标CN最优（0.305@6meas, 0.783@4meas）接近，表明膝点在整体噪声传播方面具有竞争力。Peinado等人[13]在633 nm单波长下报告的EWV约为0.35–0.40，本文膝点在350–700 nm宽波段下的EWV=0.318处于较好水平。')

add_h2(doc, '4.3 局限性')
add_body(doc, '（1）LCVR延迟量模型基于器件手册和文献数据，实际器件可能存在批次差异。温度系数的线性假设在20–50°C范围内是合理的近似[6]，但Xie等人[7]指出超过该范围后需考虑非线性的温漂行为。')
add_body(doc, '（2）仿真环境中的噪声模型假设高斯+泊松分布，未考虑1/f噪声、读出噪声等实际探测器非理想特性。Bruce等人[8]的实验误差分析表明实际系统的误差来源更为复杂。')
add_body(doc, '（3）六次测量方案需施加6组电压进行6次独立测量，测量时间相比四次测量增加50%，对动态场景可能不适用。')
add_body(doc, '（4）与Chang等人[15]的解析推导不同，本文采用纯数值优化方法，未从理论上证明最优目标函数的解析形式。两种路径各有优劣。')

# ====== 5. 结论 ======
add_h1(doc, '5 结论')
add_body(doc, '本文对双LCVR Stokes偏振测量系统的测量配置进行了多目标优化，主要结果如下：')
add_body(doc, '（1）六次测量(2×3)方案在条件数、BCPN、EWV和温漂四个指标上均优于四次测量(2×2)方案。与随机配置相比，优化后的膝点配置将Stokes重建误差降低了63%。')
add_body(doc, '（2）快轴角度在±3°误差范围内，膝点配置的条件数变化小于5%，表明六次过测量方案对装配误差具有一定的宽容性。')
add_body(doc, '（3）与已有文献[12,14,15]的对比表明，本文膝点配置的CN（≈2.245）处于或略优于现有工作的水平，EWV（0.318）和温漂（0.035）提供了额外的性能参考。')
add_body(doc, '（4）推荐配置：θ₁=−65.74°, θ₂=−41.29°, V₁=[8.44, 1.59, 2.06] V, V₂=[9.72, 1.95] V。')

# ====== 参考文献 ======
add_h1(doc, '参考文献')
refs = [
    '[1] J. S. Tyo, D. L. Goldstein, D. B. Chenault, and J. A. Shaw, "Review of passive imaging polarimetry for remote sensing applications," Applied Optics, vol. 45, no. 22, pp. 5453–5469, 2006.',
    '[2] N. Ghosh and I. A. Vitkin, "Tissue polarimetry: concepts, challenges, applications, and outlook," Journal of Biomedical Optics, vol. 16, no. 11, 110801, 2011.',
    '[3] X. Chen, S. Liu, and T. Yang, "Condition-number-based measurement configuration optimization for nanostructure reconstruction by optical scatterometry," Optics Express, vol. 31, no. 15, pp. 24466–24481, 2023.',
    '[4] J. E. Wolfe and R. A. Chipman, "Polarimetric characterization of liquid-crystal-on-silicon panels," Applied Optics, vol. 45, no. 8, pp. 1688–1703, 2006.',
    '[5] A. Peinado, A. Lizana, J. Vidal, C. Iemmi, and J. Campos, "Optimization and polarimeter design based on liquid-crystal variable retarders," Proc. SPIE, vol. 7790, 77900Q, 2010.',
    '[6] Y. Niu, L. Duan, J. Zhang, J. Huang, Y. Zhai, and W. Quan, "Suppression of ambient temperature-caused drift in a laser power stabilization system with a liquid crystal variable retarder in atomic gyroscopes," Review of Scientific Instruments, vol. 93, no. 5, 053001, 2022.',
    '[7] H. Xie, X. Zhou, Z. Cao, X. Li, and W. Zhao, "Nonlinear polarization drift modeling and real-time compensation in LCVR-based laser power stabilization systems," Applied Optics, vol. 64, no. 15, pp. 4032–4041, 2025.',
    '[8] N. C. Bruce and I. Montes-González, "The effect of experimental errors on optimized polarimeters," Proc. SPIE, vol. 12690, 126900G, 2023.',
    '[9] D. S. Sabatke, M. R. Descour, E. L. Dereniak, W. C. Sweatt, S. A. Kemme, and G. S. Phipps, "Figures of merit for complete Stokes polarimeter optimization," Applied Optics, vol. 39, no. 16, pp. 2686–2694, 2000.',
    '[10] D. Lara and C. Paterson, "Stokes polarimeter optimization in the presence of shot and Gaussian noise," Optics Express, vol. 17, no. 23, pp. 21240–21256, 2009.',
    '[11] F. Goudail, "Noise propagation and equivalent number of bits of polarimetric measurement systems," Optics Letters, vol. 38, no. 14, pp. 2458–2460, 2013.',
    '[12] P. A. Letnes, I. S. Nerbø, L. M. S. Aas, P. G. Ellingsen, and M. Kildemo, "Fast and optimal broad-band Stokes/Mueller polarimeter design by the use of a genetic algorithm," Optics Express, vol. 18, no. 22, pp. 23095–23103, 2010.',
    '[13] A. Peinado, A. Lizana, J. Vidal, C. Iemmi, A. Márquez, I. Moreno, and J. Campos, "Analysis, optimization and implementation of a variable retardance based polarimeter," EPJ Web of Conferences, vol. 5, 06008, 2010.',
    '[14] J. Wu, Y. Li, T. Ning, C. Long, and G. Zhou, "Optimal design for a broadband Stokes polarimeter of liquid crystal variable retarders," Applied Optics, vol. 61, no. 28, pp. 8254–8262, 2022.',
    '[15] L. Chang, J. Li, Y. Zhang, Y. Yin, and J. Liu, "Optimization analysis of a Stokes polarimeter for broadband liquid crystal variable retarders under the optimal objective function," JOSA A, vol. 41, no. 4, pp. 694–704, 2024.',
    '[16] W. Jiang, L. Ren, and J. Liang, "LCVR-based broadband time-division polarimetric imaging technology in visible spectral range," Optics and Lasers in Engineering, vol. 187, 109581, 2026.',
    '[17] I. Montes-González, J. Bisbal, D. Li, M. Canabal-Carbia, J. Campos, I. Estévez, and Á. Lizana, "Optimization and development of a compact NIR-RGB Mueller matrix imaging polarimeter: application to in situ skin measurements," Adv. Photon. Nexus, vol. 5, no. 2, 026009, 2026.',
    '[18] L. Chang, J. Liu, Y. Yin, R. Zhang, X. Zhao, and Z. Li, "Optimization of broadband ferroelectric liquid crystal Stokes polarimeters: evaluation and selection," Proc. SPIE, vol. 13517, 1351709, 2025.',
    '[19] K. Deb, A. Pratap, S. Agarwal, and T. Meyarivan, "A fast and elitist multiobjective genetic algorithm: NSGA-II," IEEE Trans. Evolutionary Computation, vol. 6, no. 2, pp. 182–197, 2002.',
    '[20] A. E. Eiben and J. E. Smith, Introduction to Evolutionary Computing, 2nd ed. Springer, 2015.',
    '[21] J. Yang, X. Zhang, J. Hou, and Y. Shen, "Three-dimensional instrument polarization analysis and optimization of liquid-crystal-based Stokes polarimeter," Applied Optics, vol. 62, no. 36, pp. 9587–9596, 2023.',
    '[22] C. Lai, Y. Li, J. Wu, and J. Liu, "The optimal measurement matrix with minimal condition number for double Stokes-Mueller polarimetry," Proc. SPIE, vol. 12758, 127580S, 2023.',
]
for ref in refs:
    p = doc.add_paragraph(); pf = p.paragraph_format; pf.line_spacing = 1.5
    r = p.add_run(ref); set_font(r, '宋体', 'Times New Roman', Pt(9))

save_path = os.path.expanduser('~/research-project/lcvr-optimization/paper/manuscript_zh_v3.docx')
doc.save(save_path)
print(f"V3已保存: {save_path}")
